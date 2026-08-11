from __future__ import annotations

import csv
import math
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from statistics import mean, median

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(r"C:\Users\cwb\Desktop\job\01二维公路软件\平整度验证\新算法平整度栗庙路真值验证")
OUT = Path(r"D:\job\工作COD\2025\公路\XRDataProcess\output\doc\新算法平整度LP等效修正_18工程验证总结.docx")

NAVY, BLUE, TEAL, RED = "18324A", "215A78", "2A8C82", "B42318"
LIGHT_BLUE, LIGHT_TEAL, LIGHT_RED, LIGHT_GRAY = "EAF3F7", "E8F5F2", "FBE9E7", "F3F5F7"
TEXT = "39434D"


@dataclass
class Report:
    group: str
    project: str
    side: str
    path: Path
    segments: int
    mean_error: float
    median_error: float
    p95_error: float
    max_error: float
    over1: int
    over5: int
    clamped: int
    global_rounds: str
    local_corrections: str
    target_software_max_diff: float | None
    failures: list[dict]


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "bottom", "left", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "D4DDE3")


def font(run, size=10, color=TEXT, bold=False) -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def text(p, value: str, size=10, color=TEXT, bold=False) -> None:
    run = p.add_run(str(value))
    font(run, size, color, bold)


def spacing(p, before=0, after=5, line=1.2) -> None:
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def heading(doc, value: str, level=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    spacing(p, before=12 if level == 1 else 8, after=5)
    text(p, value, 16 if level == 1 else 12.5, NAVY if level == 1 else BLUE, True)


def para(doc, value: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    spacing(p)
    if bold_prefix and value.startswith(bold_prefix):
        text(p, bold_prefix, bold=True, color=NAVY)
        text(p, value[len(bold_prefix):])
    else:
        text(p, value)


def bullet(doc, value: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    spacing(p, after=2)
    text(p, value)


def table(doc, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    header = t.rows[0]
    for i, h in enumerate(headers):
        c = header.cells[i]
        c.width = Cm(widths[i])
        shade(c, NAVY)
        border(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, 2, 2, 1.0)
        text(p, h, 8.6, "FFFFFF", True)
    for r, values in enumerate(rows):
        row = t.add_row()
        for i, value in enumerate(values):
            c = row.cells[i]
            c.width = Cm(widths[i])
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r % 2:
                shade(c, LIGHT_GRAY)
            border(c)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            spacing(p, 1, 1, 1.0)
            color = RED if ("超5" in str(value) or str(value).endswith("是")) else TEXT
            text(p, value, 8.4, color, color == RED)


def callout(doc, title: str, body: str, warning=False) -> None:
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    shade(c, LIGHT_RED if warning else LIGHT_TEAL)
    p = c.paragraphs[0]
    spacing(p, 4, 4)
    text(p, title + "  ", 10.5, RED if warning else TEAL, True)
    text(p, body, 10.2)


def page_number(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    font(run, 8, "75818A")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def pct(values: list[float], q: float) -> float:
    return sorted(values)[min(len(values) - 1, math.ceil(len(values) * q) - 1)]


def parse_report(path: Path) -> Report:
    with path.open(encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f, delimiter="\t"))
    errors = [float(row["相对误差(%)"]) for row in rows]
    iri_path = path.with_name("IRI_10m.txt")
    target_diff = None
    if iri_path.exists():
        software = [float(line.split()[1]) for line in iri_path.read_text(encoding="utf-8-sig").splitlines() if len(line.split()) >= 2]
        if software:
            target_diff = max(abs(float(row["目标IRI"]) - software[i]) for i, row in enumerate(rows[:len(software)]))
    failures = []
    for row in rows:
        if float(row["相对误差(%)"]) > 5:
            failures.append(row)
    return Report(
        group=path.parents[3].name,
        project=path.parents[2].name,
        side=path.parent.name,
        path=path,
        segments=len(rows),
        mean_error=mean(errors),
        median_error=median(errors),
        p95_error=pct(errors, 0.95),
        max_error=max(errors),
        over1=sum(e > 1 for e in errors),
        over5=sum(e > 5 for e in errors),
        clamped=sum(row.get("倍率限幅", "否") == "是" for row in rows),
        global_rounds=rows[0].get("全局修正轮数", rows[0].get("修正轮数", "未记录")),
        local_corrections=rows[0].get("局部修正次数", "未记录"),
        target_software_max_diff=target_diff,
        failures=failures,
    )


def build() -> Path:
    reports = sorted((parse_report(p) for p in BASE.rglob("LP_KB_Test_Result.txt")), key=lambda x: (x.group, x.project))
    all_segments = sum(r.segments for r in reports)
    all_errors: list[float] = []
    for r in reports:
        with r.path.open(encoding="utf-8-sig", newline="") as f:
            all_errors.extend(float(row["相对误差(%)"]) for row in csv.DictReader(f, delimiter="\t"))
    all_over5 = sum(r.over5 for r in reports)
    all_over1 = sum(r.over1 for r in reports)
    passed_reports = sum(r.over5 == 0 for r in reports)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    page_number(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n")
    text(p, "新算法平整度 LP 等效修正", 25, NAVY, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text(p, "栗庙路真值验证 - 18份工程报告汇总", 16, BLUE, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text(p, "统计时间：2026年7月30日    数据范围：30、50、70三组，共18份DAQ0报告", 9.5, TEXT)
    doc.add_paragraph("\n")
    callout(doc, "验证结论", f"18份报告共 {all_segments} 个10m区间，{all_segments-all_over5}/{all_segments}（{(all_segments-all_over5)/all_segments*100:.2f}%）在5%误差内；16/18份报告没有超过5%的区间。4个超过5%的区间集中在50组的两份工程，未出现倍率限幅。")

    heading(doc, "1. 技术摘要")
    para(doc, f"本次按当前平整度LP等效修正方法，对18份DAQ0工程报告进行汇总。整体误差均值为 {mean(all_errors):.3f}%，中位数为 {median(all_errors):.3f}%，P95为 {pct(all_errors, 0.95):.3f}%，最大值为 {max(all_errors):.3f}%。")
    bullet(doc, f"5%口径：{all_segments-all_over5}/{all_segments}段通过，整体通过率 {(all_segments-all_over5)/all_segments*100:.2f}%。")
    bullet(doc, f"1%口径：{all_segments-all_over1}/{all_segments}段通过，整体通过率 {(all_segments-all_over1)/all_segments*100:.2f}%。")
    bullet(doc, f"工程样本口径：{passed_reports}/18份报告全部满足5%；另外2份报告存在4个超过5%的10m区间。")
    bullet(doc, "所有报告的“倍率限幅”均为0段，说明异常不是0.10/5.00安全限幅造成。")

    heading(doc, "2. 分组结果：50组需要优先复核")
    groups = defaultdict(list)
    for r in reports:
        groups[r.group].append(r)
    group_rows = []
    for group, items in sorted(groups.items()):
        seg = sum(x.segments for x in items)
        group_rows.append([
            group,
            str(len(items)),
            str(seg),
            f"{mean(x.mean_error for x in items):.3f}%",
            f"{max(x.max_error for x in items):.3f}%",
            str(sum(x.over1 for x in items)),
            str(sum(x.over5 for x in items)),
        ])
    table(doc, ["分组", "报告数", "10m段数", "报告均值均值", "最大误差", ">1%段数", ">5%段数"], group_rows, [1.4, 1.5, 2.0, 3.0, 2.2, 2.1, 2.1])
    para(doc, "30组的6份报告全部在1%内；50组出现全部4个超过5%的区间；70组没有超过5%，但70-01的1%超限段较多，应作为后续精度优化的重点样本。")

    heading(doc, "3. 超过5%的区间：集中而非普遍")
    failure_rows = []
    for r in reports:
        for row in r.failures:
            failure_rows.append([
                r.project.split("__")[0],
                row["段号"],
                f"{float(row['速度(km/h)']):.3f}",
                f"{float(row['原始IRI']):.3f}",
                f"{float(row['目标IRI']):.3f}",
                f"{float(row['LP重算IRI']):.3f}",
                f"{float(row['相对误差(%)']):.3f}%",
                row["倍率限幅"],
            ])
    table(doc, ["工程", "段号", "速度", "原始IRI", "目标IRI", "LP重算IRI", "误差", "限幅"], failure_rows, [2.2, 1.2, 1.7, 2.0, 2.0, 2.2, 1.6, 1.4])
    callout(doc, "解释", "这4个区间均未触发倍率限幅，因此不能通过放宽MinScale/MaxScale解决。它们更可能与局部起伏波形、相邻10m段的状态传递及局部搜索上限有关，应使用同一工程的完整LP和IRI状态链进一步定位。", True)

    heading(doc, "4. 方法与统计口径")
    para(doc, "每个10m区间先由原始LP计算原始IRI R，再按速度档位从Coeff.dat选择K、B，得到软件目标值 T=K×R+B。LP并不直接做“高程×K+B”，而是围绕区间首尾基准坡线缩放局部起伏，再用连续四分之一车模型对整条LP复算。")
    para(doc, "报告误差定义为 |LP重算IRI - 目标IRI| / 目标IRI。车体状态在10m边界连续传递，因此单段修正会影响后续段；当前程序先进行整线同步修正，再在必要时对超限段及其前一段进行局部全线复算搜索。")
    bullet(doc, "本报告的5%与1%均为每10m段相对误差阈值。")
    bullet(doc, "“局部修正次数”是整条路线累计接受的局部候选次数，不是单段次数。")
    bullet(doc, "本报告基于内部同算法复算；客户侧验收还需保证采样间隔、矩阵、初始状态、分段和舍入位数一致。")

    heading(doc, "5. 全部18份报告明细")
    detail_rows = []
    for r in reports:
        name = r.project.split("__")[0]
        status = "通过" if r.over5 == 0 else f"超5% {r.over5}段"
        target_check = "已核对" if r.target_software_max_diff is not None and r.target_software_max_diff < 0.00001 else "待核对"
        detail_rows.append([
            r.group,
            name,
            str(r.segments),
            f"{r.mean_error:.3f}%",
            f"{r.p95_error:.3f}%",
            f"{r.max_error:.3f}%",
            str(r.over1),
            str(r.over5),
            str(r.clamped),
            str(r.global_rounds),
            str(r.local_corrections),
            status,
            target_check,
        ])
    table(doc, ["组", "工程", "段数", "均值", "P95", "最大", ">1%", ">5%", "限幅", "全局", "局部", "5%结论", "目标核对"], detail_rows, [0.8, 1.6, 0.9, 1.15, 1.15, 1.15, 0.85, 0.85, 0.85, 0.9, 0.9, 1.65, 1.4])

    heading(doc, "6. 结论与建议")
    para(doc, "总体上，当前算法在这批18份验证报告中已达到较高的5%一致性：1134个10m段中1130段通过。需继续关注的不是全局性失效，而是少数局部波形复杂的工程样本。")
    bullet(doc, "优先复核50-04和50-06：保存最终导出LP、IRI_10m.txt、Speed_10m.txt和LP_KB_Test_Result.txt，定位超5%段及其前一段的纵断面波形。")
    bullet(doc, "对70-01和70-06，以1%为优化目标复核局部搜索上限及候选策略；两者没有5%风险，但存在较多1%超限段。")
    bullet(doc, "对外交付前，应由客户实际使用的IRI程序复算最终LP；内部报告用于预检，不能替代客户算法验收。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
