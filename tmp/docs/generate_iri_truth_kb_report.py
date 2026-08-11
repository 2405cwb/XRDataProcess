import json
from collections import defaultdict
from datetime import date
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(r"C:\Users\cwb\Desktop\job\01二维公路软件\平整度验证\新算法平整度栗庙路真值验证")
DATA = Path(r"D:\job\工作COD\2025\公路\XRDataProcess\tmp\docs\iri_truth_kb_analysis.json")
OUT = Path(r"D:\job\工作COD\2025\公路\XRDataProcess\output\doc\栗庙路真值校准_分速度K_B拟合分析报告.docx")
TRUE_ORDER = ["200-300", "300-400", "400-500", "500-600"]
SLOT = {"30": "Coeff.dat 的 40 km/h 档（实际速度 25<v≤40）", "50": "Coeff.dat 的 60 km/h 档（实际速度 40<v≤60）", "70": "Coeff.dat 的 75 km/h 档（实际速度 60<v≤75）"}

def set_cell(cell, text, bold=False):
    cell.text = str(text)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Microsoft YaHei"
            r.font.size = Pt(8.5)
            r.bold = bold

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for c, h in zip(t.rows[0].cells, headers): set_cell(c, h, True)
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row): set_cell(c, v)
    if widths:
        for row in t.rows:
            for c, w in zip(row.cells, widths): c.width = Cm(w)
    doc.add_paragraph()
    return t

def metric(samples, key_pred):
    e = np.array([s[key_pred]-s["truth"] for s in samples], dtype=float)
    rel = np.abs(e)/np.array([s["truth"] for s in samples])
    return np.sqrt(np.mean(e*e)), np.mean(np.abs(e)), np.mean(rel), np.max(rel)

def main():
    x = json.loads(DATA.read_text(encoding="utf-8"))
    samples = x["samples"]
    groups = x["groups"]
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2"):
        styles[name].font.name = "Microsoft YaHei"

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("栗庙路真值校准\n分速度 K、B 拟合分析报告"); r.bold=True; r.font.name="Microsoft YaHei"; r.font.size=Pt(20)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"数据目录：{ROOT}\n编制日期：{date.today():%Y-%m-%d}").font.size=Pt(9)

    doc.add_heading("1. 结论摘要", level=1)
    doc.add_paragraph("建议采用按实际速度档选择参数的方案，推荐使用鲁棒回归（Huber）得到的新 K、B。该方案在18个工程、72个有效100m样本上拟合；留一工程交叉验证表明，相较现有参数，三个速度档的平均绝对误差均明显降低。")
    table(doc,["速度试验组", "实际速度范围 (km/h)", "建议写入档位", "新K", "新B", "留一工程 MAE", "留一工程最大绝对误差"],[
        [g, f"{m['speed_min']:.2f}–{m['speed_max']:.2f}", SLOT[g], f"{m['huber']['k']:.6f}", f"{m['huber']['b']:.6f}", f"{m['huber_cv']['mae']:.3f}", f"{m['huber_cv']['max_abs']:.3f}"]
        for g,m in groups.items()
    ], [1.8,3.0,5.0,2.1,2.1,2.7,3.2])
    doc.add_paragraph("注意：建议值只覆盖本次实际采集速度范围。15、25、130 km/h 档没有真值样本，不应据此改动。30、50、70只是试验组名称，软件实际应按 Coeff.dat 的速度阈值选档。")

    doc.add_heading("2. 输入数据与真值", level=1)
    doc.add_paragraph("读取“原始平整度结果”目录下18份 IRI_100m.xlsx：30、50、70速度组各6份。每份表格取200–300、300–400、400–500、500–600四个100m区间，合计72个样本。经用户确认，表中IRI为未乘K、未加B的原始计算结果；因此本报告直接以该列作为基础IRI，不再进行任何旧K、B反算。")
    table(doc,["里程区间 (m)", "真值 IRI"], [[a,b] for a,b in [("200–300","1.25"),("300–400","1.27"),("400–500","1.43"),("500–600","1.75")]], [5,4])

    doc.add_heading("3. 数学方法", level=1)
    doc.add_paragraph("设原始表中未乘K、未加B的100m平整度为 I_raw。按实际车速将样本分入三个速度档，直接对每个速度档独立拟合线性关系：")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("I_true ≈ K_new × I_raw + B_new").bold=True
    doc.add_paragraph("拟合采用 Huber 鲁棒回归：小残差按平方误差处理，大残差按线性损失降低其影响。报告同时计算普通最小二乘（OLS）供对照。为检验不是只记住某一个工程，采用留一工程交叉验证：每次拿出一个工程，用其余5个工程拟合，再预测被拿出的4个真值区间。")

    doc.add_heading("4. 参数与误差对比", level=1)
    rows=[]
    for g,m in groups.items():
        ss=[s for s in samples if s['group']==g]
        old_rmse,old_mae,old_rel,old_max=metric(ss,'current_iri')
        new_rmse,new_mae,new_rel,new_max=metric(ss,'huber_pred')
        rows.append([g, f"{old_mae:.3f}",f"{new_mae:.3f}",f"{(1-new_mae/old_mae)*100:.1f}%",f"{old_rel*100:.1f}%",f"{new_rel*100:.1f}%",f"{new_max*100:.1f}%"])
    table(doc,["速度组", "原始IRI MAE", "新参数 MAE", "MAE降低", "原始IRI平均相对误差", "新参数平均相对误差", "新参数样本最大相对误差"],rows,[1.6,2,2,2,2.8,3,3])
    table(doc,["速度组", "OLS K", "OLS B", "OLS R²", "Huber K（建议）", "Huber B（建议）", "留一工程CV RMSE"],[
        [g,f"{m['ols']['k']:.6f}",f"{m['ols']['b']:.6f}",f"{m['ols']['r2']:.3f}",f"{m['huber']['k']:.6f}",f"{m['huber']['b']:.6f}",f"{m['huber_cv']['rmse']:.3f}"] for g,m in groups.items()
    ],[1.5,2,2,1.5,2.7,2.7,2.7])
    doc.add_paragraph("选择Huber值的原因：它与OLS的拟合效果接近，但在50、70组的留一工程最大误差更小或相当，且对偶发异常值更稳健。")

    doc.add_heading("5. 建议的 Coeff.dat 修改范围", level=1)
    doc.add_paragraph("由于输入表本身未应用K、B，本次得到的是应施加到原始IRI上的参数。建议在测试副本中将下列三个速度阈值行写入新K、B；其余行不改。")
    table(doc,["Coeff.dat速度阈值", "对应本次样本", "输入表已应用K", "输入表已应用B", "建议新K", "建议新B"],[
        ["40", "30组，25<v≤40", "否", "否", f"{groups['30']['huber']['k']:.6f}", f"{groups['30']['huber']['b']:.6f}"],
        ["60", "50组，40<v≤60", "否", "否", f"{groups['50']['huber']['k']:.6f}", f"{groups['50']['huber']['b']:.6f}"],
        ["75", "70组，60<v≤75", "否", "否", f"{groups['70']['huber']['k']:.6f}", f"{groups['70']['huber']['b']:.6f}"],
    ],[3,4,2,2,2.5,2.5])

    doc.add_heading("6. 使用边界与下一步验证", level=1)
    doc.add_paragraph("这是一套“100m IRI输出校准”的经验参数，不等同于对LP高程直接加B。若目标是让客户用导出的LP重算IRI后也与软件结果一致，仍需沿用已实现的LP等效修正流程，并用完整线路重算验证。")
    doc.add_paragraph("四个真值区间来自同一路段，虽然在18个工程和三个速度范围内重复验证，但其路面谱特征有限。建议至少再补充另一条路、每档不少于4个独立100m真值区间；验收时以新工程的绝对误差和相对误差共同评价。当前数据上，个别留出工程仍可能超过5%相对误差，不能承诺所有单段必然小于5%。")

    doc.add_heading("附录A：72个样本明细（Huber建议参数预测）", level=1)
    landscape=doc.add_section(WD_SECTION.NEW_PAGE); landscape.orientation=1
    landscape.page_width,landscape.page_height=landscape.page_height,landscape.page_width
    landscape.left_margin=landscape.right_margin=Cm(1.2)
    detail=[]
    for g in ('30','50','70'):
        for s in sorted([z for z in samples if z['group']==g],key=lambda z:(z['project'],TRUE_ORDER.index(z['interval']))):
            detail.append([g,s['project'].split('__')[0],s['interval'],f"{s['speed']:.2f}",f"{s['raw_iri']:.3f}",f"{s['current_iri']:.3f}",f"{s['truth']:.2f}",f"{s['huber_pred']:.3f}",f"{s['huber_error']:+.3f}"])
    table(doc,["组","工程","区间m","速度","原始IRI","输入表IRI","真值","新参数预测","预测-真值"],detail,[1,3,2,1.8,2.5,2.3,1.5,2.3,2.3])
    OUT.parent.mkdir(parents=True,exist_ok=True)
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__': main()
