from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\job\工作COD\2025\公路\XRDataProcess")
OUT_DIR = ROOT / "output" / "doc"
OUT_FILE = OUT_DIR / "IRI_KB等效反映到LP方案说明.docx"

NAVY = "18324A"
BLUE = "215A78"
TEAL = "2A8C82"
LIGHT_BLUE = "EAF3F7"
LIGHT_TEAL = "E8F5F2"
LIGHT_GRAY = "F2F4F6"
DARK_GRAY = "3B4650"
WHITE = "FFFFFF"
AMBER = "B36B00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, east_asia="微软雅黑", latin="Aptos", size=10.5,
                 color=DARK_GRAY, bold=False, italic=False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=17, color=NAVY, bold=True)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), TEAL)
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        set_run_font(run, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=NAVY)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_paragraph_spacing(p, after=3)
    add_text(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=3)
    add_text(p, text)
    return p


def add_equation(doc, equation, note=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "18", "color": TEAL},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=4)
    add_text(p, equation, east_asia="等线", latin="Cambria Math",
             size=12, color=NAVY, bold=True)
    if note:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, after=6)
        add_text(p2, note, size=9, color=DARK_GRAY)
    return table


def add_callout(doc, title, body, kind="info"):
    fill = LIGHT_TEAL if kind == "info" else "FFF4E5"
    accent = TEAL if kind == "info" else AMBER
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "22", "color": accent},
        top={"val": "single", "sz": "4", "color": fill},
        right={"val": "single", "sz": "4", "color": fill},
        bottom={"val": "single", "sz": "4", "color": fill},
    )
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=4, after=2)
    add_text(p, title + "  ", size=10.5, color=accent, bold=True)
    add_text(p, body, size=10.5, color=DARK_GRAY)
    return table


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=3, after=3, line=1.0)
        add_text(p, text, size=9.5, color=WHITE, bold=True)
        if widths:
            cell.width = Cm(widths[i])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            if row_index % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=2, after=2, line=1.1)
            add_text(p, str(value), size=9.2)
            if widths:
                cell.width = Cm(widths[i])
    return table


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=8.5, color="7A858E")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    for style_name in ("Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10.5)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "技术方案说明", size=13, color=TEAL, bold=True)
    p.paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10)
    add_text(p, "IRI 的 K、B 校正", size=28, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=18)
    add_text(p, "等效反映到国检 LP 的数学逻辑与实现方案",
             size=19, color=BLUE, bold=True)
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(line, "━━━━━━━━━━━━━━━━━━━━━━━━━━", size=12, color=TEAL)
    line.paragraph_format.space_after = Pt(22)
    add_callout(
        doc,
        "核心目标",
        "客户无需知道内部 K、B 参数，只使用导出的 LP 和标准 IRI 算法，"
        "即可复算出与软件校正后 IRI 尽可能一致的结果。",
    )
    doc.add_paragraph("\n\n\n")
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta_data = [
        ("项目", "XRDataProcess 国检转换"),
        ("文档性质", "算法设计、实现映射与验证说明"),
        ("版本", "V1.0"),
        ("日期", "2026年7月24日"),
    ]
    for r, (k, v) in enumerate(meta_data):
        meta.cell(r, 0).width = Cm(3.5)
        meta.cell(r, 1).width = Cm(10.5)
        set_cell_shading(meta.cell(r, 0), LIGHT_BLUE)
        for c, text in enumerate((k, v)):
            p = meta.cell(r, c).paragraphs[0]
            set_paragraph_spacing(p, before=3, after=3)
            add_text(p, text, size=10, color=NAVY if c == 0 else DARK_GRAY,
                     bold=(c == 0))
            set_cell_border(
                meta.cell(r, c),
                top={"val": "single", "sz": "4", "color": "D7E0E5"},
                bottom={"val": "single", "sz": "4", "color": "D7E0E5"},
                left={"val": "single", "sz": "4", "color": "D7E0E5"},
                right={"val": "single", "sz": "4", "color": "D7E0E5"},
            )
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    add_page_number(section)

    add_cover(doc)

    add_heading(doc, "1. 结论先行", 1)
    add_body(
        doc,
        "这套方案是可行的，但不能把结果公式中的 B 直接加到每个高程点上。"
        "正确方向是：先把软件希望得到的校正后 IRI 转换成每个 10 m 区间的"
        "“等效起伏倍率”，再围绕该区间的基准坡线放大或缩小纵断面起伏，"
        "最终用完整 IRI 算法复算和校验。"
    )
    add_callout(
        doc,
        "一句话概括",
        "K、B 是 IRI 结果层的修正参数；LP 是高程输入层的数据。"
        "两者之间不存在简单的逐点加法，需要通过“目标 IRI -> 起伏倍率 -> LP -> IRI 复算”完成反演。",
    )
    add_bullet(doc, "软件内部目标：每 10 m 的校正结果为 Tj = Kj × Rj + Bj。")
    add_bullet(doc, "客户侧输入：只有校正后的 LP，不提供速度档位、K、B。")
    add_bullet(doc, "客户侧输出：用约定的标准算法计算 Aj，并要求 Aj 与 Tj 的相对误差尽量小。")
    add_bullet(doc, "当前代码验收目标设为整线最大相对误差不超过 1%，最多接受 5 轮改进。")

    add_heading(doc, "2. 问题定义", 1)
    add_heading(doc, "2.1 符号", 2)
    add_table(
        doc,
        ["符号", "含义"],
        [
            ("j", "第 j 个 10 m 区间"),
            ("h", "原始 LP 的整条纵断面高程序列"),
            ("Fj(h)", "标准 IRI 算法在连续状态下对第 j 段的计算结果"),
            ("Rj", "原始 LP 计算出的第 j 段原始 IRI，即 Rj = Fj(h)"),
            ("vj", "第 j 段车速"),
            ("Kj、Bj", "按 vj 从 Coeff.dat 对应速度档位选出的修正参数"),
            ("Tj", "软件希望客户最终复算得到的目标 IRI"),
            ("sj", "第 j 段纵断面起伏的等效倍率"),
            ("h*", "最终导出的等效校正 LP"),
            ("Aj", "客户使用 h* 复算得到的第 j 段 IRI"),
        ],
        widths=[3.0, 12.0],
    )

    add_heading(doc, "2.2 软件目标与客户目标", 2)
    add_equation(doc, "Rj = Fj(h)")
    add_equation(doc, "Tj = Kj × Rj + Bj",
                 "Kj、Bj 由该 10 m 区间的速度 vj 选择。")
    add_equation(doc, "Aj = Fj(h*)  ≈  Tj",
                 "最终交付目标：客户只读取 h*，无需知道 K、B。")

    add_heading(doc, "3. 为什么不能直接对 LP 高程做 K × h + B", 1)
    add_heading(doc, "3.1 高程整体加 B 对 IRI 基本不起作用", 2)
    add_body(
        doc,
        "IRI 的激励来源是相邻采样点之间的高程变化。若所有高程同时加上同一个常数 B，"
        "相邻点高差不会改变："
    )
    add_equation(doc, "(hi+1 + B) - (hi + B) = hi+1 - hi")
    add_body(
        doc,
        "因此，结果公式中的“+B”不能靠给 LP 的每个点统一增加 B 来实现。"
        "这只会整体抬高纵断面，不会产生需要的附加不平整度。"
    )
    add_heading(doc, "3.2 整条高程直接乘 K 会同时改变纵坡", 2)
    add_body(
        doc,
        "若使用 h'i = K × hi，会把绝对高程、长坡趋势和局部起伏一起缩放。"
        "这样虽然可能改变 IRI，但也会破坏 LP 的整体几何意义，而且仍不能正确表达结果层的 B。"
    )
    add_callout(
        doc,
        "关键区别",
        "应调整的是“相对于局部基准坡线的起伏”，而不是绝对高程。",
        kind="warning",
    )

    add_heading(doc, "4. 一次等效变换的数学构造", 1)
    add_heading(doc, "4.1 建立每 10 m 的基准坡线", 2)
    add_body(
        doc,
        "对第 j 个 10 m 区间，以该段首点和末点连成直线。对区间内第 i 个采样点，"
        "基准高程写作 Lj,i："
    )
    add_equation(
        doc,
        "Lj,i = hstart + (hend - hstart) × (i - istart) / (iend - istart)"
    )
    add_body(doc, "原始高程相对基准线的局部起伏为：")
    add_equation(doc, "qj,i = hj,i - Lj,i")
    add_heading(doc, "4.2 将 K、B 转换为起伏倍率", 2)
    add_body(
        doc,
        "在 IRI 对局部幅值近似线性的前提下，第一轮可用目标值与原始值的比例"
        "作为等效起伏倍率："
    )
    add_equation(doc, "sj(0) = Tj / Rj = Kj + Bj / Rj")
    add_body(
        doc,
        "这个式子解释了 B 如何进入 LP：B 没有作为高程常数相加，"
        "而是通过 Bj/Rj 转换成与该段原始粗糙程度相关的附加倍率。"
    )
    add_heading(doc, "4.3 重建等效 LP", 2)
    add_equation(doc, "h*j,i = Lj,i + sj × (hj,i - Lj,i)")
    add_table(
        doc,
        ["性质", "结果"],
        [
            ("端点保持", "区间首尾点的起伏为 0，因此首尾高程不变"),
            ("纵坡保持", "首尾连线代表的区间总体纵坡不变"),
            ("起伏可控", "sj > 1 放大局部起伏；0 < sj < 1 缩小局部起伏"),
            ("B 的表达", "B 通过 Bj/Rj 进入 sj，而不是直接加到高程上"),
        ],
        widths=[3.5, 11.5],
    )
    add_heading(doc, "4.4 从数学上看，这是一个低维反问题", 2)
    add_body(
        doc,
        "原始 LP 有数万个高程点，直接逐点反求既不唯一也不稳定。本方案不把每个高程点"
        "都作为未知数，而是把每个 10 m 的起伏倍率 sj 作为未知数。这样，315 段路线只需要"
        "求 315 个倍率，且每个倍率都有明确的物理含义和限幅条件。"
    )
    add_equation(doc, "h*(s) = Baseline(h) + Scale(s) × Fluctuation(h)")
    add_equation(
        doc,
        "minimize  E(s) = maxj { |Fj(h*(s)) - Tj| / Tj }",
        "约束：0.10 ≤ sj ≤ 5.00；最终高程按 LP 实际精度舍入。"
    )
    add_body(
        doc,
        "初值 sj(0)=Tj/Rj 来自“IRI 对整体起伏幅值近似成比例”的一阶近似。"
        "随后不求解析导数，而是通过完整算法复算得到真实误差，再用阻尼比例更新。"
        "因此它本质上是带边界约束、以整线最大误差为目标的数值反演。"
    )
    add_callout(
        doc,
        "适用边界",
        "当 Rj 接近 0 时，Bj/Rj 会过大。当前实现用 MinRawIri 判断并对倍率限制在 [0.10, 5.00]；"
        "极平顺路段仍需单独关注。",
        kind="warning",
    )

    add_heading(doc, "5. 为什么一次变换后仍需要整线复算", 1)
    add_body(
        doc,
        "四分之一车模型具有内部状态。第 j 段开始时的车身和悬架状态来自前一段末尾，"
        "10 m 边界只重置该段的累计量，不会把车辆状态清零。因此第 j 段的 IRI 不仅受"
        "本段高程影响，也会受到前面路面激励留下的状态影响。"
    )
    add_equation(doc, "xj,start = xj-1,end")
    add_body(
        doc,
        "这意味着各段并不是互相独立的。修改某一段起伏后，相邻段乃至后续段的复算结果"
        "可能一起变化；高程舍入、末段长度和客户算法实现差异也会产生误差。"
    )
    add_callout(
        doc,
        "回答常见疑问",
        "不是“前一个 IRI 数值”直接参与后一个 IRI，而是前面高程序列形成的四分之一车内部状态"
        "被带入下一段。",
    )

    add_heading(doc, "6. 误差压缩与安全迭代", 1)
    add_heading(doc, "6.1 完整复算", 2)
    add_body(
        doc,
        "用第一轮倍率重建整条 LP，按最终导出的小数位数舍入，然后从路线起点连续运行"
        "与客户约定一致的 IRI 算法，得到 Aj。"
    )
    add_equation(doc, "ej = |Aj - Tj| / Tj")
    add_equation(doc, "E = max(e1, e2, ..., en)",
                 "E 是整条路线所有 10 m 区间中的最大相对误差。")
    add_heading(doc, "6.2 候选倍率更新", 2)
    add_equation(doc, "s'j = sj × (Tj / Aj)^α", "α 为步长，初值 1.0。")
    add_body(
        doc,
        "每次候选都从原始 LP 重新构造，不在上一轮已经修改的高程上继续叠加，"
        "从而避免累计变形。倍率仍限制在 [0.10, 5.00]。"
    )
    add_heading(doc, "6.3 接受、回退和步长减半", 2)
    add_number(doc, "用候选倍率重建整条 LP，并按实际导出精度舍入。")
    add_number(doc, "连续复算整条路线，求候选最大误差 E'。")
    add_number(doc, "只有 E' < E 时才接受候选；否则丢弃候选并将 α 减半。")
    add_number(doc, "步长依次尝试 1、0.5、0.25、0.125、0.0625、0.03125。")
    add_number(doc, "没有任何步长能降低最大误差时停止，保留上一轮已经验证更好的 LP。")
    add_callout(
        doc,
        "关于收敛",
        "不能证明任意数据都一定收敛，也不能事先承诺必然达到某个误差。"
        "能够保证的是：代码只接受最大误差严格下降的候选，所以每个已接受轮次的 E 单调下降；"
        "若无法继续改善则停止。",
        kind="warning",
    )

    add_heading(doc, "7. 完整流程", 1)
    add_table(
        doc,
        ["步骤", "输入", "处理", "输出"],
        [
            ("1", "原始 LP", "用连续四分之一车模型按 10 m 计算", "Rj"),
            ("2", "速度 vj、Coeff.dat", "按速度档位选择同一组参数", "Kj、Bj"),
            ("3", "Rj、Kj、Bj", "计算 Tj = Kj×Rj+Bj", "目标 IRI"),
            ("4", "Tj、Rj", "计算初始 sj = Tj/Rj，并限幅", "初始倍率"),
            ("5", "原始 LP、sj", "缩放相对基准线的起伏并舍入", "候选 LP"),
            ("6", "候选 LP", "连续复算整条路线", "Aj、最大误差 E"),
            ("7", "Tj、Aj", "阻尼更新、整线复算、接受或回退", "改进倍率"),
            ("8", "满足停止条件的 LP", "写出国检 LP 和测试报告", "客户交付文件"),
        ],
        widths=[1.2, 3.5, 7.0, 3.3],
    )

    add_heading(doc, "8. 停止条件与可调参数", 1)
    add_table(
        doc,
        ["参数", "当前值", "作用", "风险提示"],
        [
            ("TargetRelativeError", "0.01（1%）", "整线最大相对误差达到该值即停止", "调得更小会增加迭代，但不保证一定达到"),
            ("MaxCorrectionIterations", "5", "最多接受的改进轮数", "增加轮数只提供更多机会，不等于保证更准"),
            ("MinScale / MaxScale", "0.10 / 5.00", "限制起伏倍率", "防止 Rj 很小时出现异常放大"),
            ("最小步长 α", "0.03125", "候选失败时逐级减半", "再小可能收益低且增加计算量"),
            ("LP 舍入位数", "sheetRoundingOffNum", "按客户实际读取文件精度验证", "双方精度不一致会引入复算差异"),
        ],
        widths=[4.0, 2.7, 5.2, 4.0],
    )
    add_body(
        doc,
        "若希望把误差进一步压小，优先调整 TargetRelativeError，其次适量增加 "
        "MaxCorrectionIterations；但必须以重新导出的 LP 报告为准，不能只看循环次数。"
    )

    add_heading(doc, "9. 当前代码对应关系", 1)
    add_table(
        doc,
        ["代码位置", "当前职责"],
        [
            ("GlobalExcel.cs:6295", "TargetRelativeError = 0.01，当前目标为 1%"),
            ("GlobalExcel.cs:6296", "MaxCorrectionIterations = 5"),
            ("GlobalExcel.cs:6326", "SelectLpKbBySpeed：按速度选择配对的 K、B"),
            ("GlobalExcel.cs:6347", "BuildCorrectedLpProfile：首次构造等效 LP"),
            ("GlobalExcel.cs:6358-6416", "整线复算、阻尼候选、最大误差比较、接受或回退"),
            ("GlobalExcel.cs:6472", "按每 10 m 基准坡线缩放局部起伏"),
            ("GlobalExcel.cs:6521", "CalculateMaxLpIriError：计算整线最大误差"),
            ("MyIRIMTD_new.cs:1910", "软件 IRI 结果执行 irival = irival × kparm + bparm"),
            ("MyIRIMTD_new.cs:1606-1653", "四分之一车 oldZSU 状态沿路线连续传递"),
        ],
        widths=[5.2, 10.8],
    )
    add_callout(
        doc,
        "避免重复校正",
        "LP 等效校正应从原始纵断面构造。若在 LP 上先做旧的平均 K 乘法，再执行本方案，"
        "会发生重复校正。",
        kind="warning",
    )

    add_heading(doc, "10. 实际工程验证", 1)
    add_heading(doc, "10.1 验证工程与文件", 2)
    add_body(
        doc,
        "验证工程：第二次_低速路_上行_1_湖北省_武汉市_江夏区_20250724_143214。"
        "本次结果生成于 2026年7月24日 11:27，导出 LP 长度约 3.142 km，"
        "包含 DAQ0、DAQ1 两侧数据。验证使用以下文件："
    )
    add_bullet(doc, r"IRIMTD\DAQ0、DAQ1\Coeff.dat：速度档位及 K、B。")
    add_bullet(doc, r"IRIMTD\DAQ0、DAQ1\IRI_10m.txt：软件校正后 IRI。")
    add_bullet(doc, r"IRIMTD\DAQ0、DAQ1\LP_KB_Test_Result.txt：等效 LP 内部复算报告。")
    add_bullet(doc, r"ConverSource\X000999999A\RIFile\X000999999A-LP-0.000-20250724143214.txt：最终导出 LP。")

    add_heading(doc, "10.2 内部等效目标的复算结果", 2)
    add_body(
        doc,
        "报告中的“目标 IRI”按 Tj=Kj×Rj+Bj 生成，“LP 重算 IRI”由最终舍入后的等效 LP"
        "连续复算。两侧各 315 段，所有区间对该内部目标均满足 1%："
    )
    add_table(
        doc,
        ["指标", "DAQ0", "DAQ1"],
        [
            ("区间数", "315", "315"),
            ("报告修正轮数", "2", "5"),
            ("平均相对误差", "0.0765%", "0.0105%"),
            ("中位相对误差", "0.033%", "0.000%"),
            ("P95 相对误差", "0.316%", "0.021%"),
            ("最大相对误差", "0.840%", "0.850%"),
            ("误差不超过 1%", "315/315", "315/315"),
            ("倍率范围", "0.868393 - 1.109057", "0.849676 - 1.035753"),
            ("倍率触发限幅", "0 段", "0 段"),
        ],
        widths=[6.0, 4.5, 4.5],
    )
    add_body(
        doc,
        "该结果表明：对于当前内部目标函数，阻尼迭代确实把两侧整线最大误差压到了 1% 以内。"
        "DAQ0 在较少轮次内达到阈值；DAQ1 使用了更多轮次，说明两侧波形和状态耦合程度不同，"
        "不能预设固定一轮就足够。"
    )

    add_heading(doc, "10.3 与软件 IRI_10m.txt 的逐段对照", 2)
    add_body(
        doc,
        "进一步把 LP 重算结果直接与软件输出 IRI_10m.txt 对照，前 314 个完整 10 m 区间"
        "的目标值与软件值仅存在约 10^-7 量级的文本舍入差异。完整 10 m 区间的结果如下："
    )
    add_table(
        doc,
        ["指标（前 314 个完整 10 m 区间）", "DAQ0", "DAQ1"],
        [
            ("平均相对误差", "0.0760%", "0.0106%"),
            ("最大相对误差", "0.8404%", "0.8499%"),
            ("误差不超过 1%", "314/314", "314/314"),
        ],
        widths=[7.0, 4.0, 4.0],
    )

    add_heading(doc, "10.4 末段约 2 m 的边界问题", 2)
    add_body(
        doc,
        "导出 LP 的末桩号为 3.142000 km，因此第 315 段不是完整 10 m，而是约 2 m 的残余段。"
        "该段在 LP 校正过程与软件 IRI 输出过程中的采样范围或末段归一化口径尚未完全对齐。"
        "直接相对软件 IRI 比较如下："
    )
    add_table(
        doc,
        ["侧别", "软件 IRI", "LP 重算 IRI", "相对误差", "5% 判定"],
        [
            ("DAQ0", "4.361399", "4.498021", "3.1325%", "通过"),
            ("DAQ1", "2.385086", "2.556491", "7.1865%", "未通过"),
        ],
        widths=[2.5, 3.2, 3.2, 3.0, 2.6],
    )
    add_callout(
        doc,
        "验证结论",
        "若只统计完整 10 m 区间，两侧 628 个结果全部在 1% 内；若把约 2 m 的末段也按软件值验收，"
        "DAQ1 当前超过客户常用的 5% 要求。报告时必须单列末段，不能表述为全线所有区间均已满足 1%。",
        kind="warning",
    )
    add_body(
        doc,
        "下一步应统一 CalculateLpIri10M 与 GenerateIRI_NEW 对末段的有效采样起点、终点、count、"
        "partial_distance 和速度取值，再重新导出验证。工程长度恰好是 10 m 整数倍时不会暴露"
        "这一类残余段口径差异，但通用交付仍必须修正。"
    )

    add_heading(doc, "11. 对外使用条件与风险", 1)
    add_bullet(
        doc,
        "客户算法必须与软件验证算法在采样间隔、四分之一车矩阵、初始状态、"
        "10 m 分段和末段处理上保持一致。"
    )
    add_bullet(
        doc,
        "本方案实现的是“对指定 IRI 算法的数值等效”，不代表修改后的 LP 仍等同于"
        "原始物理测量纵断面；建议对外明确文件用途。"
    )
    add_bullet(
        doc,
        "速度档位突变会导致相邻 10 m 使用不同倍率，虽然端点保持连续，"
        "仍应检查边界附近的波形和后续状态影响。"
    )
    add_bullet(
        doc,
        "极低 Rj、短末段和倍率触发限幅的区间应单独列入报告，"
        "不能用整体合格率掩盖局部异常。"
    )
    add_bullet(
        doc,
        "若客户实现与内部算法不完全一致，应以客户真实程序对最终 LP 的复算结果"
        "作为最终验收依据。"
    )

    add_heading(doc, "12. 验收建议", 1)
    add_table(
        doc,
        ["验收层级", "建议标准"],
        [
            ("内部算法一致性", "最终写盘 LP 由内部同算法复算，整线最大误差达到设定阈值；未达到时明确列出"),
            ("客户算法一致性", "客户程序对同一 LP 独立复算，所有 10 m 区间误差不超过合同或项目约定"),
            ("文件一致性", "确认客户实际收到的 LP 与内部验证文件哈希一致"),
            ("异常审计", "报告倍率限幅、迭代停止原因、末段长度和所有超限区间"),
        ],
        widths=[4.0, 12.0],
    )
    add_callout(
        doc,
        "最终判定",
        "只有“客户真实算法 + 最终写盘 LP”的复算结果，才能证明对外交付误差满足要求。"
        "内部迭代达到 1% 是重要的预检条件，但不能替代客户侧验收。",
    )

    add_heading(doc, "附录 A：算法伪代码", 1)
    pseudo = [
        "rawIRI = CalculateIRI(originalLP)",
        "for each 10m section j:",
        "    (Kj, Bj) = SelectBySpeed(vj)",
        "    target[j] = max(0, Kj * rawIRI[j] + Bj)",
        "    scale[j] = clamp(target[j] / rawIRI[j], 0.10, 5.00)",
        "",
        "correctedLP = BuildFromOriginal(originalLP, scale)",
        "RoundAsExported(correctedLP)",
        "actual = CalculateIRI(correctedLP)        // 整线连续计算",
        "E = MaxError(target, actual)",
        "",
        "acceptedIterations = 0",
        "while E > 1% and acceptedIterations < 5:",
        "    alpha = 1.0",
        "    accepted = false",
        "    while alpha >= 0.03125:",
        "        candidateScale[j] = clamp(",
        "            scale[j] * pow(target[j] / actual[j], alpha),",
        "            0.10, 5.00)",
        "        candidateLP = BuildFromOriginal(originalLP, candidateScale)",
        "        RoundAsExported(candidateLP)",
        "        candidateActual = CalculateIRI(candidateLP)",
        "        candidateE = MaxError(target, candidateActual)",
        "        if candidateE < E:",
        "            accept candidate; accepted = true; break",
        "        alpha = alpha / 2",
        "    if not accepted: break",
        "",
        "write correctedLP and LP_KB_Test_Result.txt",
    ]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FA")
    set_cell_border(
        cell,
        top={"val": "single", "sz": "4", "color": "CBD5DB"},
        bottom={"val": "single", "sz": "4", "color": "CBD5DB"},
        left={"val": "single", "sz": "4", "color": "CBD5DB"},
        right={"val": "single", "sz": "4", "color": "CBD5DB"},
    )
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=6, after=6, line=1.0)
    for idx, line in enumerate(pseudo):
        run = p.add_run(line)
        set_run_font(run, east_asia="等线", latin="Consolas", size=8.7,
                     color="263238")
        if idx != len(pseudo) - 1:
            run.add_break()

    add_heading(doc, "附录 B：方案边界", 1)
    add_body(
        doc,
        "本方案不是在数学上解析求得唯一 LP，而是在保持每段首尾基准坡线的约束下，"
        "寻找一组局部起伏倍率，使指定 IRI 算法的输出接近目标值。由于 IRI 只描述路面"
        "纵断面对标准车辆模型的响应，不同波形可能得到相同或相近的 IRI，所以等效 LP"
        "不是原始 LP 的唯一反解。"
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_document())
